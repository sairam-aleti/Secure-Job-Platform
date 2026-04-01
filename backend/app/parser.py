from pypdf import PdfReader
import io
import json
import logging

logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_bytes):
    """Extracts raw text from a PDF file."""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        text = ""
        for page in reader.pages:
            content = page.extract_text()
            if content:
                text += content
        return text.lower()
    except Exception as e:
        logger.warning("Text extraction failed for uploaded PDF")
        return ""

def extract_text_from_docx(file_bytes):
    """
    Extracts text from a DOCX file including paragraphs and tables.
    Uses python-docx for comprehensive extraction.
    """
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        
        parts = []
        
        # Extract all paragraph text
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)
        
        # Extract all table content
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    parts.append(" | ".join(row_text))
        
        return " ".join(parts).lower()
    except Exception as e:
        logger.warning(f"DOCX text extraction failed: {e}")
        return ""

def calculate_match_score(resume_text, job_skills_data):
    """
    Compares resume text against job skills.
    Handles both JSON strings and Python lists.
    """
    if not resume_text or not job_skills_data:
        return 0
    
    resume_text = resume_text.lower()
    
    # Handle different data types for skills
    if isinstance(job_skills_data, str):
        try:
            job_skills = json.loads(job_skills_data)
        except Exception:
            # Fallback: if it's a comma-separated string
            job_skills = [s.strip() for s in job_skills_data.split(',')]
    else:
        job_skills = job_skills_data

    if not job_skills:
        return 0
    
    match_count = 0
    for skill in job_skills:
        if skill.lower() in resume_text:
            match_count += 1
            
    score = int((match_count / len(job_skills)) * 100)
    logger.info(f"Match score calculated: {match_count}/{len(job_skills)} = {score}%")
    return score